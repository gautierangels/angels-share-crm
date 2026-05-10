import streamlit as st
import pandas as pd
from database import get_db
from utils import fmt_date, fmt_money
from datetime import date, timedelta
from pathlib import Path
import io

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

RAPPORTS_DIR = Path("/Users/gautiersalinier/Documents/Angels Share Marketing Limited/Angels' Share Management/Rapports")
RAPPORTS_DIR.mkdir(parents=True, exist_ok=True)


def _periode_label(debut, fin):
    mois = ["Janvier","Février","Mars","Avril","Mai","Juin",
            "Juillet","Août","Septembre","Octobre","Novembre","Décembre"]
    if debut.month == fin.month and debut.year == fin.year:
        return f"{mois[debut.month-1]} {debut.year}"
    if debut.month == 1 and fin.month == 3:
        return f"T1 {debut.year}"
    if debut.month == 4 and fin.month == 6:
        return f"T2 {debut.year}"
    if debut.month == 7 and fin.month == 9:
        return f"T3 {debut.year}"
    if debut.month == 10 and fin.month == 12:
        return f"T4 {debut.year}"
    return f"{fmt_date(str(debut))} — {fmt_date(str(fin))}"


def _generer_rapport_docx(producteur, commandes, interactions, periode_label,
                           include_commandes, include_interactions,
                           include_distribution, distribution_data,
                           message_perso):
    """Génère le rapport Word en mémoire et retourne les bytes."""
    doc = Document()

    # Styles généraux
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── En-tête avec logo ────────────────────────────────────────────────────
    import os
    from pathlib import Path as _Path

    # Logo Angels' Share si disponible
    logo_candidates = [
        _Path(__file__).parent.parent / "assets" / "logo.png",
        _Path(__file__).parent.parent / "assets" / "logo.jpg",
        _Path(__file__).parent.parent / "static" / "logo.png",
    ]
    logo_path = next((p for p in logo_candidates if p.exists()), None)

    if logo_path:
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(str(logo_path), width=Inches(1.8))
        except Exception:
            pass

    # Titre principal
    h = doc.add_heading("", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = h.add_run(f"Rapport d'activité — {producteur['nom']}")
    run_h.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)
    run_h.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = sub.add_run("Angels' Share Marketing Limited  |  ")
    r1.font.size = Pt(11)
    r1.font.bold = True
    r2 = sub.add_run(periode_label)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    doc.add_paragraph()

    # ── Message personnalisé ──────────────────────────────────────────────────
    if message_perso:
        doc.add_heading('Message', level=2)
        doc.add_paragraph(message_perso)
        doc.add_paragraph()

    # ── Résumé chiffres clés ──────────────────────────────────────────────────
    doc.add_heading('Résumé de la période', level=2)

    ca_total = sum(c["montant"] or 0 for c in commandes)
    nb_cmd   = len(commandes)
    clients  = list(dict.fromkeys(c["client_nom"] or "—" for c in commandes))

    resume = doc.add_table(rows=3, cols=2)
    resume.style = 'Table Grid'
    resume.alignment = WD_TABLE_ALIGNMENT.CENTER
    data_resume = [
        ("Nombre de commandes", str(nb_cmd)),
        ("Chiffre d'affaires total", fmt_money(ca_total)),
        ("Clients actifs", ", ".join(clients[:5]) + ("…" if len(clients)>5 else "")),
    ]
    for i, (label, valeur) in enumerate(data_resume):
        resume.cell(i,0).text = label
        resume.cell(i,0).paragraphs[0].runs[0].bold = True
        resume.cell(i,1).text = valeur

    doc.add_paragraph()

    # ── Commandes ─────────────────────────────────────────────────────────────
    if include_commandes and commandes:
        doc.add_heading('Commandes', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ["Proforma", "Client", "Pays", "Montant", "Statut"]
        for i, h_txt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h_txt
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            # Fond bleu header via shading
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1A5C8A')
            tc_pr.append(shd)

        for cmd in commandes:
            row = table.add_row().cells
            row[0].text = cmd["proforma"] or "—"
            row[1].text = cmd["client_nom"] or "—"
            row[2].text = cmd["pays"] or "—"
            row[3].text = fmt_money(cmd["montant"] or 0, cmd["devise"] or "EUR")
            row[4].text = cmd["statut"] or "—"

        doc.add_paragraph()

    # ── Interactions / activités ───────────────────────────────────────────────
    if include_interactions and interactions:
        doc.add_heading('Activités & Interactions', level=2)
        for inter in interactions:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{fmt_date(str(inter['date_interaction'])[:10])} — ").bold = True
            p.add_run(f"{inter['type'] or ''} : {inter['description'] or ''}")

        doc.add_paragraph()

    # ── Distribution ───────────────────────────────────────────────────────────
    if include_distribution and distribution_data:
        doc.add_heading('État de la distribution', level=2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        for i, h_txt in enumerate(["Pays", "Statut", "Client actuel"]):
            tbl.rows[0].cells[i].text = h_txt
            tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for d in distribution_data:
            row = tbl.add_row().cells
            row[0].text = d["pays"] or "—"
            row[1].text = d["statut"] or "—"
            row[2].text = d["client_actuel"] or "—"

        doc.add_paragraph()

    # ── Pied de page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"Angels' Share Marketing Limited — Rapport généré le {fmt_date(str(date.today()))}"
    )
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer_run.font.size = Pt(9)

    # Sauvegarder en bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render():
    st.markdown("## 📋 Rapports producteurs")
    db = get_db()

    producteurs = db.execute(
        "SELECT * FROM producteurs WHERE archived=0 ORDER BY nom"
    ).fetchall()

    if not producteurs:
        st.info("Aucun producteur enregistré.")
        db.close()
        return

    tab1, tab2 = st.tabs(["📄 Générer un rapport", "📁 Rapports enregistrés"])

    # ══ GÉNÉRER ═══════════════════════════════════════════════════════════════
    with tab1:
        st.markdown("### Paramètres du rapport")

        # Producteur
        prod_noms = [p["nom"] for p in producteurs]
        prod_choix = st.selectbox("🍇 Producteur", prod_noms, key="rp_prod")
        prod = next(p for p in producteurs if p["nom"] == prod_choix)

        # Période
        c1, c2, c3 = st.columns(3)
        type_periode = c1.selectbox("Période", [
            "Mois en cours", "Mois précédent",
            "Trimestre en cours", "Trimestre précédent",
            "Année en cours", "Personnalisée"
        ], key="rp_periode")

        today = date.today()
        if type_periode == "Mois en cours":
            debut = today.replace(day=1)
            fin   = today
        elif type_periode == "Mois précédent":
            fin   = today.replace(day=1) - timedelta(days=1)
            debut = fin.replace(day=1)
        elif type_periode == "Trimestre en cours":
            q = (today.month - 1) // 3
            debut = date(today.year, q*3+1, 1)
            fin   = today
        elif type_periode == "Trimestre précédent":
            q = (today.month - 1) // 3
            if q == 0: debut = date(today.year-1, 10, 1); fin = date(today.year-1, 12, 31)
            else: debut = date(today.year, (q-1)*3+1, 1); fin = date(today.year, q*3, 1) - timedelta(days=1)
        elif type_periode == "Année en cours":
            debut = date(today.year, 1, 1)
            fin   = today
        else:
            debut = c2.date_input("Du", value=today.replace(day=1), key="rp_debut")
            fin   = c3.date_input("Au", value=today, key="rp_fin")

        periode_label = _periode_label(debut, fin)
        st.caption(f"📅 Période : **{periode_label}** ({debut} → {fin})")

        st.markdown("---")

        # Options contenu
        st.markdown("### Contenu du rapport")
        o1, o2, o3 = st.columns(3)
        inc_cmd   = o1.checkbox("📦 Commandes", value=True)
        inc_inter = o2.checkbox("💬 Interactions", value=True)
        inc_dist  = o3.checkbox("🌐 Distribution", value=True)

        # Message personnalisé
        message_perso = st.text_area(
            "✏️ Message personnalisé (introduction libre)",
            placeholder="Ex: Cher Jean-Marc, voici le résumé de notre activité pour votre maison sur la période écoulée…",
            height=120,
            key="rp_message"
        )

        st.markdown("---")

        # ── Prévisualisation des données ──────────────────────────────────────
        debut_str = str(debut); fin_str = str(fin)

        commandes = db.execute("""
            SELECT * FROM commandes
            WHERE producteur_id = ?
            AND date_commande BETWEEN ? AND ?
            AND archived = 0
            ORDER BY date_commande DESC
        """, (prod["id"], debut_str, fin_str)).fetchall()

        interactions = db.execute("""
            SELECT * FROM interactions
            WHERE entite_id = ?
            AND entite_type = 'producteur'
            AND date_interaction BETWEEN ? AND ?
            ORDER BY date_interaction DESC
        """, (prod["id"], debut_str, fin_str)).fetchall()

        distribution_data = []
        if inc_dist:
            try:
                distribution_data = db.execute("""
                    SELECT pays, statut, client_actuel, marque_nom
                    FROM distribution
                    WHERE producteur_id = ?
                    AND (archived = 0 OR archived IS NULL)
                    ORDER BY pays
                """, (prod["id"],)).fetchall()
            except Exception:
                pass

        # KPIs prévisualisation
        ca_total = sum(c["montant"] or 0 for c in commandes)
        st.markdown(f"### Aperçu — {prod_choix} · {periode_label}")

        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Commandes", len(commandes))
        k2.metric("CA total", fmt_money(ca_total))
        k3.metric("Interactions", len(interactions))
        k4.metric("Pays distribués", len(distribution_data))

        if commandes:
            st.markdown("**Commandes de la période :**")
            df_cmd = pd.DataFrame([{
                "Proforma":  c["proforma"],
                "Client":    c["client_nom"] or "—",
                "Pays":      c["pays"] or "—",
                "Montant":   fmt_money(c["montant"] or 0, c["devise"] or "EUR"),
                "Statut":    c["statut"] or "—",
            } for c in commandes])
            st.dataframe(df_cmd, use_container_width=True, hide_index=True)
        else:
            st.info("Aucune commande sur cette période.")

        if interactions:
            st.markdown("**Interactions de la période :**")
            for inter in interactions[:5]:
                st.markdown(
                    f"- `{fmt_date(str(inter['date_interaction'])[:10])}` "
                    f"**{inter['type'] or ''}** — {inter['description'] or ''}"
                )
            if len(interactions) > 5:
                st.caption(f"… et {len(interactions)-5} autre(s)")

        st.markdown("---")

        # ── Génération ────────────────────────────────────────────────────────
        col_gen, col_dl = st.columns(2)

        if not DOCX_OK:
            st.warning("⚠️ `python-docx` non installé.")
            st.code("pip3 install python-docx")
            # Export texte de secours
            if col_gen.button("📝 Rapport texte", use_container_width=True):
                nom_prod = prod["nom"]
                lignes = [
                    "RAPPORT D'ACTIVITE — " + nom_prod,
                    "Periode : " + periode_label,
                    "Genere le : " + fmt_date(str(date.today())),
                    "="*60,
                    "Commandes : " + str(len(commandes)) + " | CA : " + fmt_money(ca_total),
                    "="*60,
                ]
                if message_perso:
                    lignes = ["Message : " + message_perso, ""] + lignes
                for cmd in commandes:
                    lignes.append("- " + str(fmt_date(str(cmd["date_commande"])[:10]))
                                  + " — " + str(cmd["client_nom"] or "—")
                                  + " — " + fmt_money(cmd["montant"] or 0, cmd["devise"] or "EUR"))
                txt = "\n".join(lignes)
                st.download_button("⬇️ Télécharger (.txt)", txt,
                    file_name="Rapport_" + nom_prod + "_" + periode_label + ".txt")
        else:
            if col_gen.button("📄 Générer le rapport Word", use_container_width=True, type="primary"):
                with st.spinner("Génération en cours…"):
                    try:
                        docx_bytes = _generer_rapport_docx(
                            producteur=prod,
                            commandes=commandes,
                            interactions=interactions,
                            periode_label=periode_label,
                            include_commandes=inc_cmd,
                            include_interactions=inc_inter,
                            include_distribution=inc_dist,
                            distribution_data=distribution_data,
                            message_perso=message_perso,
                        )

                        # Format : Rapport_Angels'Share_NomProducteur_Periode.docx
                        nom_prod_clean = prod["nom"].replace(" ","_").replace("/","_")
                        periode_clean  = periode_label.replace(" ","_")
                        nom_fichier = f"Rapport_AngelsShare_{nom_prod_clean}_{periode_clean}.docx"

                        # Sauvegarder localement
                        save_path = RAPPORTS_DIR / nom_fichier
                        with open(save_path, "wb") as f:
                            f.write(docx_bytes)

                        st.success(f"✅ Rapport généré : **{nom_fichier}**")

                        # Bouton téléchargement
                        st.download_button(
                            label="⬇️ Télécharger le rapport",
                            data=docx_bytes,
                            file_name=nom_fichier,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )

                    except Exception as e:
                        st.error(f"Erreur lors de la génération : {e}")

    # ══ RAPPORTS ENREGISTRÉS ══════════════════════════════════════════════════
    with tab2:
        st.markdown("### Rapports générés")
        if RAPPORTS_DIR.exists():
            fichiers = sorted(RAPPORTS_DIR.glob("*.docx"), reverse=True)
            if not fichiers:
                st.info("Aucun rapport généré pour l'instant.")
            else:
                for f in fichiers:
                    c1, c2, c3 = st.columns([4, 1, 1])
                    c1.markdown(f"📄 **{f.name}**")
                    c2.caption(f"{f.stat().st_size // 1024} Ko")
                    with open(f, "rb") as fp:
                        c3.download_button(
                            "⬇️",
                            data=fp.read(),
                            file_name=f.name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_{f.name}",
                        )
        else:
            st.info("Dossier rapports non trouvé.")

    db.close()
